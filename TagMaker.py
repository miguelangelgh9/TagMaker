import requests
from bs4 import BeautifulSoup
from docx import Document

def getTopo(s):
    temp=''
    for c in s:
        if c:
            temp=temp+c+" "
    return temp
lista=[15106280,15106287]
tags=[]
for i in lista:
    barcode=i
    r = requests.get('http://sb.ues.edu.sv/cgi-bin/koha/opac-search.pl', params={'q': barcode})
    reg=r.url.split("=")[1]

    soup=BeautifulSoup(r.text,"html.parser")

    titulo=soup.find("h1",class_="title").text.split("/")[0]
    autor=soup.find("h5",class_="author").text[4:-9]

    topo=soup.find("td",class_="call_no").text.replace("\n","").split(" ")
    topo=getTopo(topo).split(" ")
    codT=str(topo[0])
    codC=str(topo[1])
    cutter=str(topo[2])
    fecha=str(topo[3])
    tags=tags+[{"titulo":titulo,"autor":autor,"barcode":barcode,"reg":reg,"codT":codT,"codC":codC,"cutter":cutter,"fecha":fecha}]

#print(tags)
document=Document()
for i in tags:
    rec=str(i['codT'])+"\n"+str(i["codC"])+"\n"+str(i["cutter"])+"\n"+str(i["fecha"])+"\n"+str(i["barcode"])+"\n"+str(i["autor"])+"\n"+str(i["titulo"])+"\n"+"MFN:"+str(i["reg"])+"\n"
    #print(i["codC"])
    document.add_paragraph(rec)
    document.add_paragraph(
    i["codT"]+"\n"+
    i["codC"]+"\n"+
    i["cutter"]+"\n"+
    i["fecha"]+"\n"
    )
    document.add_paragraph(
    i["autor"]+"\n"+
    i["titulo"]+"\n"+
    i["codT"]+" "+str(i["codC"])+" "+i["cutter"]+" "+str(i["fecha"])+" "+str(i["barcode"])+" "+"("+str(i["reg"])+")"
    )
    document.add_page_break()
document.save("Tags.docx")
